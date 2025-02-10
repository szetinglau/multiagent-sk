from docx import Document

# Create Internal Document
internal_doc = Document()
internal_doc.add_heading('Internal Review Document for Claim HURTX-2024-0456', level=1)
internal_doc.add_paragraph('Claim Summary:')
internal_doc.add_paragraph('Claimant: Contoso Electronics Inc.')
internal_doc.add_paragraph('Incident Type: Hurricane (Hurricane Alicia)')
internal_doc.add_paragraph('Date of Incident: May 15, 2024')
internal_doc.add_paragraph('Risk Assessment Metrics:')
internal_doc.add_picture('risk_assessment_metrics.png')
internal_doc.add_paragraph('Final Recommendation: Partial Approval')
internal_doc.add_paragraph('Approved Amount: $480,000')
internal_doc.add_paragraph('Conditions: Further documentation required for business interruption claims.')
internal_doc.save('report/HURTX-2024-0456_Internal.docx')